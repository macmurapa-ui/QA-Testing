#!/usr/bin/env python3
"""
generate_coverage.py
====================
Scans all JS test files in this repo for @covers tags, then regenerates
QA_COVERAGE.md and QA_COVERAGE.docx from coverage_areas.json.

Tag format (add near the top of any JS test file):
    // @covers 2.1, 4.1

Run manually:  python3 generate_coverage.py
Runs auto:     pre-commit hook (.git/hooks/pre-commit)
"""

import json
import os
import re
import glob
from datetime import date
from collections import OrderedDict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ──────────────────────────────────────────────────────────────────
GREEN     = RGBColor(0x1E, 0x8E, 0x3E)
ORANGE    = RGBColor(0xF5, 0x7C, 0x00)
GREY      = RGBColor(0x5F, 0x6B, 0x7A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK      = RGBColor(0x1A, 0x1A, 0x2E)
HEADER_BG = RGBColor(0x1A, 0x1A, 0x2E)
ALT_BG    = RGBColor(0xF3, 0xF4, 0xF6)

STATUS_COLOUR = {'AUTOMATED': GREEN, 'PENDING': GREY, 'BLOCKED': ORANGE}

AREA_ORDER = [
    'Admin Portal',
    'Admin Reporting',
    'Agent Dashboard',
    'Partner API',
    'Billing Preferences',
]

SUBAREA_ORDER = {
    'Admin Portal':  ['Authentication & Session', 'Branch Management', 'User Management',
                      'Landlord Management', 'Property Management', 'Submission Management'],
    'Partner API':   ['Branches', 'Landlords', 'Properties', 'Moves'],
}

SPRINT_RECS = [
    ('Sprint 1 — Complete the Core Create Flow (High ROI)',
     ['Move-out submission (6.2) — highest business value; directly triggers OVO switch revenue',
      'Move-in submission (6.3) — closes the void period loop',
      'View submissions list (6.1) — foundational for all reporting tests']),
    ('Sprint 2 — Dashboard Automation via Impersonation',
     ['Dashboard access via impersonation (8.1) — prerequisite for all Dashboard tests',
      'Move-out wizard Dashboard (8.5) — agent-led submission; primary revenue driver',
      'Move-in wizard Dashboard (8.6) — void closure']),
    ('Sprint 3 — Reporting + Data Integrity',
     ['Conversions report (7.1) — validate OVO outcome codes (SALI, DONS, FDFD)',
      'CWN status (7.5) — council notification compliance',
      'Erroneous transfer check (7.6) — regulatory red line']),
    ('Sprint 4 — Partner API Coverage',
     ['POST branch / landlord / property / move (9.1, 9.4, 9.8, 9.13, 9.14) — core create flow via API',
      'GET / PATCH / DELETE (9.2, 9.3, 9.5-9.7, 9.9-9.12, 9.15-9.18) — full 18-endpoint sweep']),
    ('Sprint 5 — Billing Preferences',
     ['Billing preference V2 (10.2) — when V2 development ships; 61 scenario matrix ready']),
]

NOTES = [
    '**Auth**: Google CAPTCHA prevents fully headless OAuth. Workaround: user provides fresh session token from browser DevTools; written to `auth.json`.',
    '**Lambda caching**: DB query endpoint caches in-memory. Tautology workaround partially effective. Graeme (CTO) owns the Lambda — prompt raised.',
    '**Branch naming**: All test branches follow `Mac N[DDMMYY]` (e.g. `Mac 1200526`). `nextBranchName()` in `htm_clone_007.js` handles sequential counting.',
    '**Soft deletes**: 10 tables use `acts_as_paranoid` — always assert `deleted_at IS NULL` in DB queries.',
    '**Money**: All commission values stored in pence. Divide by 100 for GBP display.',
]

REPO_ROOT = os.path.dirname(os.path.abspath(__file__))


# ── Data loading & scanning ───────────────────────────────────────────────────

def load_areas():
    with open(os.path.join(REPO_ROOT, 'coverage_areas.json'), encoding='utf-8') as f:
        return json.load(f)


def scan_js_files(areas):
    """Find @covers tags in JS files and update area status/script."""
    area_map = {a['id']: a for a in areas}
    pattern = os.path.join(REPO_ROOT, '**', '*.js')
    for js_path in sorted(glob.glob(pattern, recursive=True)):
        try:
            with open(js_path, encoding='utf-8', errors='ignore') as f:
                content = f.read()
        except OSError:
            continue
        for match in re.finditer(r'//\s*@covers\s+([\d.,\s]+)', content):
            raw = match.group(1)
            ids = [x.strip() for x in re.split(r'[,\s]+', raw)
                   if x.strip() and re.match(r'^\d+\.\d+$', x.strip())]
            script_name = os.path.relpath(js_path, REPO_ROOT)
            for area_id in ids:
                if area_id in area_map:
                    area_map[area_id]['status'] = 'AUTOMATED'
                    area_map[area_id]['script'] = os.path.basename(js_path)
    return areas


def stats(areas):
    automated = sum(1 for a in areas if a['status'] == 'AUTOMATED')
    blocked   = sum(1 for a in areas if a['status'] == 'BLOCKED')
    pending   = sum(1 for a in areas if a['status'] == 'PENDING')
    total     = len(areas)
    pct       = round(automated / total * 100) if total else 0
    return automated, pending, blocked, total, pct


def group_areas(areas):
    """Return OrderedDict: area -> (sub_area | None) -> [area_dicts]"""
    grouped = OrderedDict()
    for area_name in AREA_ORDER:
        grouped[area_name] = OrderedDict()
    for a in areas:
        area = a['area']
        sub  = a.get('sub_area')
        if area not in grouped:
            grouped[area] = OrderedDict()
        if sub not in grouped[area]:
            grouped[area][sub] = []
        grouped[area][sub].append(a)
    # Sort sub_areas
    for area_name, subs in grouped.items():
        order = SUBAREA_ORDER.get(area_name)
        if order:
            sorted_subs = OrderedDict()
            for key in order:
                if key in subs:
                    sorted_subs[key] = subs[key]
            for key, val in subs.items():
                if key not in sorted_subs:
                    sorted_subs[key] = val
            grouped[area_name] = sorted_subs
    return grouped


# ── Markdown generation ───────────────────────────────────────────────────────

def status_icon(status):
    return {'AUTOMATED': '✅ AUTOMATED', 'PENDING': '🔲 PENDING', 'BLOCKED': '⏳ BLOCKED'}.get(status, status)


def md_table(rows, headers):
    lines = []
    lines.append('| ' + ' | '.join(headers) + ' |')
    lines.append('|' + '|'.join(['---'] * len(headers)) + '|')
    for row in rows:
        lines.append('| ' + ' | '.join(str(c) for c in row) + ' |')
    return '\n'.join(lines)


def generate_md(areas, path, automated, pending, blocked, total, pct):
    grouped = group_areas(areas)
    lines = []
    today = date.today().strftime('%d %b %Y')

    lines.append('# HTM QA Automation Coverage\n')
    lines.append(f'**Last updated:** {today}  ')
    lines.append('**QA Test Lead:** Mac Murapa  ')
    lines.append('**Test environment:** HTM Clone (`admin-clone.helpthemove.co.uk`) | Agent Dashboard | Partner API\n')
    lines.append('---\n')
    lines.append('## Overall Coverage\n')
    lines.append(md_table(
        [['Total functional test areas defined', total],
         ['Automated', automated],
         ['Pending', pending],
         ['Blocked', blocked],
         ['**Coverage**', f'**~{pct}%**']],
        ['Metric', 'Value']
    ))

    area_counts = {'Admin Portal': 26, 'Admin Reporting': 7, 'Agent Dashboard': 8,
                   'Partner API': 18, 'Billing Preferences': 2}

    for area_name, subs in grouped.items():
        count = area_counts.get(area_name, sum(len(v) for v in subs.values()))
        lines.append(f'\n---\n\n## {area_name} ({count} areas)\n')

        if area_name == 'Partner API':
            lines.append('Docs: `https://api.helpthemove.co.uk/docs/partner` | Auth: Bearer token per clone environment\n')

        for sub_name, sub_areas in subs.items():
            if sub_name:
                lines.append(f'\n### {sub_name}\n')

            has_script_col = area_name not in ('Partner API',)
            if has_script_col:
                headers = ['#', 'Test Area', 'Status', 'Script', 'Notes']
                rows = [[a['id'], a['name'], status_icon(a['status']),
                         f"`{a['script']}`" if a['script'] else '—', a.get('notes', '')]
                        for a in sub_areas]
            else:
                headers = ['#', 'Endpoint', 'Status', 'Notes']
                rows = [[a['id'], f"`{a['name']}`", status_icon(a['status']), a.get('notes', '')]
                        for a in sub_areas]
            lines.append(md_table(rows, headers))

    lines.append('\n---\n\n## Priority Recommendations — Next Test Sprints\n')
    for sprint_name, items in SPRINT_RECS:
        lines.append(f'\n### {sprint_name}\n')
        for idx, item in enumerate(items, 1):
            lines.append(f'{idx}. **{item.split(" — ")[0]}** — {item.split(" — ", 1)[1] if " — " in item else ""}')

    lines.append('\n---\n\n## Notes\n')
    for note in NOTES:
        lines.append(f'- {note}')

    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines) + '\n')


# ── DOCX generation ───────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_col = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def docx_table(doc, headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.add_row()
    for j, h in enumerate(headers):
        cell = hdr_row.cells[j]
        set_cell_bg(cell, HEADER_BG)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
    for i, row_data in enumerate(rows):
        row = table.add_row()
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_cell_border(cell)
            if i % 2 == 0:
                set_cell_bg(cell, ALT_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = re.sub(r'[✅🔲⏳]', '', str(val)).strip()
            run = p.add_run(text)
            run.font.size = Pt(9)
            for key, colour in STATUS_COLOUR.items():
                if key in str(val).upper():
                    run.font.color.rgb = colour
                    run.bold = True
                    break
    return table


def generate_docx(areas, path, automated, pending, blocked, total, pct):
    grouped = group_areas(areas)
    doc = Document()
    today = date.today().strftime('%d %b %Y')

    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    title = doc.add_heading('HTM QA Automation Coverage', level=0)
    title.runs[0].font.color.rgb = DARK

    p = doc.add_paragraph()
    p.add_run('Last updated: ').bold = False
    p.add_run(today).bold = True
    p.add_run('   QA Test Lead: ').bold = False
    p.add_run('Mac Murapa').bold = True

    doc.add_heading('Overall Coverage', level=1)
    docx_table(doc, ['Metric', 'Value'], [
        ['Total functional test areas defined', str(total)],
        ['Automated', str(automated)],
        ['Pending', str(pending)],
        ['Blocked', str(blocked)],
        ['Coverage', f'~{pct}%'],
    ])

    area_counts = {'Admin Portal': 26, 'Admin Reporting': 7, 'Agent Dashboard': 8,
                   'Partner API': 18, 'Billing Preferences': 2}

    for area_name, subs in grouped.items():
        count = area_counts.get(area_name, sum(len(v) for v in subs.values()))
        doc.add_heading(f'{area_name} ({count} areas)', level=1)

        if area_name == 'Partner API':
            doc.add_paragraph('Docs: https://api.helpthemove.co.uk/docs/partner  |  Auth: Bearer token per clone environment')

        for sub_name, sub_areas in subs.items():
            if sub_name:
                doc.add_heading(sub_name, level=2)

            has_script = area_name not in ('Partner API',)
            if has_script:
                headers = ['#', 'Test Area', 'Status', 'Script', 'Notes']
                rows = [[a['id'], a['name'], status_icon(a['status']),
                         a['script'] or '—', a.get('notes', '')]
                        for a in sub_areas]
            else:
                headers = ['#', 'Endpoint', 'Status', 'Notes']
                rows = [[a['id'], a['name'], status_icon(a['status']), a.get('notes', '')]
                        for a in sub_areas]
            docx_table(doc, headers, rows)
            doc.add_paragraph()

    doc.add_heading('Priority Recommendations — Next Test Sprints', level=1)
    for sprint_name, items in SPRINT_RECS:
        doc.add_heading(sprint_name, level=2)
        for item in items:
            p = doc.add_paragraph(style='List Number')
            parts = item.split(' — ', 1)
            r = p.add_run(parts[0])
            r.bold = True
            if len(parts) > 1:
                p.add_run(f' — {parts[1]}')

    doc.add_heading('Notes', level=1)
    for note in NOTES:
        p = doc.add_paragraph(style='List Bullet')
        clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', note)
        clean = re.sub(r'`([^`]+)`', r'\1', clean)
        p.add_run(clean)

    doc.save(path)


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    areas = load_areas()
    areas = scan_js_files(areas)

    automated, pending, blocked, total, pct = stats(areas)

    md_path   = os.path.join(REPO_ROOT, 'QA_COVERAGE.md')
    docx_path = os.path.join(REPO_ROOT, 'QA_COVERAGE.docx')

    generate_md(areas, md_path, automated, pending, blocked, total, pct)
    generate_docx(areas, docx_path, automated, pending, blocked, total, pct)

    print(f'Coverage updated: {automated}/{total} ({pct}%) automated')
    print(f'  → {md_path}')
    print(f'  → {docx_path}')


if __name__ == '__main__':
    main()
